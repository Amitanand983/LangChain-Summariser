import validators
import streamlit as st
from langchain.prompts import PromptTemplate
from langchain_groq import ChatGroq
from langchain.chains.summarize import load_summarize_chain
from langchain.chains import RetrievalQA
from langchain.schema import Document
from langchain.vectorstores import FAISS
from langchain.embeddings import HuggingFaceEmbeddings
from langchain_community.document_loaders import YoutubeLoader, UnstructuredURLLoader
from youtube_transcript_api import YouTubeTranscriptApi
from youtube_transcript_api._errors import TranscriptsDisabled, NoTranscriptFound
import whisper
from PyPDF2 import PdfReader
import docx
import tempfile
import os
import pickle

# Streamlit App Configuration
st.set_page_config(
    page_title="LangChain: Summarize and Chat",
    page_icon="🦜",
)
st.title("🦜 LangChain: Summarize and Chat")
st.subheader("Summarize URL, Audio, or Document and Chat with the Data")

# Sidebar for Groq API Key Input
with st.sidebar:
    groq_api_key = st.text_input("Groq API Key", value="", type="password")

# Inputs
generic_url = st.text_input("URL", label_visibility="collapsed")
audio_file = st.file_uploader("Upload an audio file (mp3, wav, m4a)", type=["mp3", "wav", "m4a"])
document_file = st.file_uploader("Upload a document file (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])

# LangChain LLM Setup (Gemma Model)
llm = ChatGroq(model="Gemma-7b-It", groq_api_key=groq_api_key)

# Prompt Templates
map_prompt_template = """
Write a concise summary of the following:
"{text}"
CONCISE SUMMARY:
"""
map_prompt = PromptTemplate(template=map_prompt_template, input_variables=["text"])

combine_prompt_template = """
Write a concise summary of the following summaries:
"{text}"
CONCISE SUMMARY:
"""
combine_prompt = PromptTemplate(template=combine_prompt_template, input_variables=["text"])

qa_prompt_template = """
Use the following pieces of context to answer the question at the end. 
If you don't know the answer, just say that you don't know, don't try to make up an answer.

Context: {context}

Question: {question}

Answer the question thoroughly and accurately based on the context provided.
"""
qa_prompt = PromptTemplate(
    template=qa_prompt_template,
    input_variables=["context", "question"]
)

# Token limits
CHUNK_SIZE = 500  # Smaller chunks for better context retrieval
CHUNK_OVERLAP = 50

def extract_text_from_document(file):
    """Extract text from uploaded document files."""
    if file.name.endswith(".pdf"):
        pdf_reader = PdfReader(file)
        return " ".join(page.extract_text() for page in pdf_reader.pages)
    elif file.name.endswith(".docx"):
        doc = docx.Document(file)
        return " ".join(paragraph.text for paragraph in doc.paragraphs)
    elif file.name.endswith(".txt"):
        return file.read().decode("utf-8")
    else:
        return ""

def process_content(content, source):
    """Process content into chunks and create Document objects."""
    # Split content into smaller chunks for better retrieval
    words = content.split()
    chunks = []
    for i in range(0, len(words), CHUNK_SIZE):
        chunk = " ".join(words[i:i + CHUNK_SIZE + CHUNK_OVERLAP])
        chunks.append(Document(page_content=chunk, metadata={"source": source}))
    return chunks

# Initialize session states
if "vectorstore" not in st.session_state:
    st.session_state.vectorstore = None
if "summarized_text" not in st.session_state:
    st.session_state.summarized_text = ""
if "original_docs" not in st.session_state:
    st.session_state.original_docs = []

# Embeddings initialization
embeddings = HuggingFaceEmbeddings()

# Button for Processing Content
if st.button("Process Content"):
    if not groq_api_key.strip() or (not generic_url.strip() and not audio_file and not document_file):
        st.error("Please provide valid input and Groq API Key.")
    else:
        try:
            with st.spinner("Processing content..."):
                docs = []

                # Process YouTube or Website URL
                if generic_url:
                    if "youtube.com" in generic_url or "youtu.be" in generic_url:
                        try:
                            video_id = (
                                generic_url.split("v=")[-1].split("&")[0]
                                if "youtube.com" in generic_url
                                else generic_url.split("/")[-1]
                            )
                            transcript = YouTubeTranscriptApi.get_transcript(video_id)
                            content = " ".join([t["text"] for t in transcript])
                            docs.extend(process_content(content, generic_url))
                        except TranscriptsDisabled:
                            st.error("Transcripts are disabled for this video.")
                        except Exception as e:
                            st.error(f"Failed to fetch YouTube content: {e}")
                    else:
                        try:
                            loader = UnstructuredURLLoader(urls=[generic_url])
                            unstructured_docs = loader.load()
                            for doc in unstructured_docs:
                                docs.extend(process_content(doc.page_content, generic_url))
                        except Exception as e:
                            st.error(f"Failed to load URL content: {e}")

                # Process Audio File
                if audio_file:
                    with tempfile.NamedTemporaryFile(delete=True) as temp_file:
                        temp_file.write(audio_file.read())
                        temp_file.flush()
                        model = whisper.load_model("base")
                        transcription = model.transcribe(temp_file.name)
                        audio_content = transcription["text"]
                        docs.extend(process_content(audio_content, "Uploaded Audio"))

                # Process Document File
                if document_file:
                    doc_content = extract_text_from_document(document_file)
                    docs.extend(process_content(doc_content, document_file.name))

                if not docs:
                    st.error("No content found. Please check the input and try again.")
                else:
                    # Store original documents and create vectorstore
                    st.session_state.original_docs = docs
                    st.session_state.vectorstore = FAISS.from_documents(docs, embeddings)
                    
                    # Create a summary using map_reduce with proper prompts
                    summarize_chain = load_summarize_chain(
                        llm,
                        chain_type="map_reduce",
                        map_prompt=map_prompt,
                        combine_prompt=combine_prompt,
                        verbose=True
                    )
                    st.session_state.summarized_text = summarize_chain.run(docs)
                    
                    st.success("Content processed successfully!")
                    st.subheader("Summary")
                    st.write(st.session_state.summarized_text)

        except Exception as e:
            st.error(f"An error occurred: {e}")
            st.exception(e)

# Chatbot Interface
if st.session_state.vectorstore:
    st.subheader("Chat with the Data")
    user_question = st.text_input("Ask a question about the content:")
    if user_question:
        with st.spinner("Thinking..."):
            # Create QA chain with custom prompt
            qa_chain = RetrievalQA.from_chain_type(
                llm=llm,
                chain_type="stuff",
                retriever=st.session_state.vectorstore.as_retriever(
                    search_kwargs={"k": 3}  # Retrieve more relevant chunks
                ),
                chain_type_kwargs={
                    "prompt": qa_prompt
                },
                return_source_documents=True
            )
            
            # Get the answer
            result = qa_chain({"query": user_question})
            answer = result["result"]
            sources = result.get("source_documents", [])

            # Display the answer and sources
            st.write("### Answer:")
            st.write(answer)
            
            if sources:
                st.write("### Sources:")
                unique_sources = set()
                for source in sources:
                    source_name = source.metadata.get('source', 'Unknown Source')
                    if source_name not in unique_sources:
                        unique_sources.add(source_name)
                        st.write(f"- {source_name}")
                        