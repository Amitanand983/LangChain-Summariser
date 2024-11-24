import validators
import streamlit as st
from langchain.prompts import PromptTemplate
from langchain_groq import ChatGroq
from langchain.chains.summarize import load_summarize_chain
from langchain.chains import RetrievalQA
from langchain.schema import Document
from langchain.vectorstores import FAISS
from langchain.embeddings import HuggingFaceEmbeddings
from langchain_community.document_loaders import UnstructuredURLLoader
from youtube_transcript_api import YouTubeTranscriptApi
from youtube_transcript_api._errors import TranscriptsDisabled, NoTranscriptFound
import whisper
from PyPDF2 import PdfReader
import docx
import tempfile
import os
from deep_translator import GoogleTranslator
import yt_dlp
from pytube import YouTube

# Model Configuration Class
class ModelConfig:
    def __init__(self):
        self.models = {
            "youtube": {
                "provider": "groq",
                "model": "mixtral-8x7b-32768",
                "description": "Optimized for video transcript analysis"
            },
            "audio": {
                "provider": "groq",
                "model": "gemma-7b-it",
                "description": "Specialized for audio content analysis"
            },
            "url": {
                "provider": "groq",
                "model": "mixtral-8x7b-32768",  # Updated for URL processing
                "description": "Optimized for web content analysis"
            },
            "document": {
                "provider": "groq",
                "model": "llama3-8b-8192",  # Updated for document processing
                "description": "Optimized for document summarization and analysis"
            }
        }
    
    def get_model(self, content_type, api_key):
        """Get the appropriate LLM based on content type."""
        model_config = self.models.get(content_type.lower(), self.models["document"])
        
        if model_config["provider"] == "groq":
            return ChatGroq(
                model=model_config["model"],
                groq_api_key=api_key
            )
        return None

# Streamlit App Configuration
st.set_page_config(
    page_title="SummAI: Summarize and Chat",
    page_icon="❄️",
)
st.title("❄️ SummAI: Summarize and Chat")
st.subheader("Summarize URL, Audio, or Document and Chat with the Data")

# Sidebar for Groq API Key Input
with st.sidebar:
    groq_api_key = st.text_input("Groq API Key", value="", type="password")
    
    # Model Information Display
    st.subheader("🤖 Model Information")
    model_config = ModelConfig()
    for content_type, model_info in model_config.models.items():
        with st.expander(f"{content_type.title()} Model"):
            st.write(f"**Model**: {model_info['model']}")
            st.write(f"**Provider**: {model_info['provider'].title()}")
            st.write(f"**Description**: {model_info['description']}")

    # History Section
    st.subheader("📜 History")
    if "history" in st.session_state and st.session_state.history:
        for idx, item in enumerate(st.session_state.history):
            st.markdown(f"**{idx + 1}. {item['type']}**: {item['name']}")
            if st.button(f"View {item['type']} {idx + 1}", key=f"view_history_{idx}"):
                st.session_state.current_summary = item["summary"]
                st.session_state.current_qa_pairs = item["qa_pairs"]

# Prompt Templates
map_prompt_template = """
Break down and analyze the following content into a structured format:
"{text}"

Please organize the summary in the following format:
1. Main Topics:
   - List key topics/themes
   - Each point should be concise but informative

2. Subtopics and Details:
   - Break down each main topic into relevant subtopics
   - Include important details and explanations

3. Key Examples/Evidence:
   - List specific examples, case studies, or evidence mentioned
   - Include relevant data points or statistics if present

4. Important Takeaways:
   - List practical implications or conclusions
   - Highlight significant findings or results

Please ensure all points are clear, concise, and well-organized with proper bullet points.
Avoid paragraph format and focus on structured, point-wise information.

STRUCTURED SUMMARY:
"""
map_prompt = PromptTemplate(template=map_prompt_template, input_variables=["text"])

combine_prompt_template = """
Combine the following summaries into a single, well-organized structure:
"{text}"

Create a comprehensive summary in this format:
1. Main Topics & Themes:
   - Consolidate all major topics
   - List overarching themes

2. Detailed Breakdown:
   - Key concepts and ideas
   - Important definitions
   - Critical arguments or positions

3. Supporting Evidence:
   - Notable examples
   - Statistical data
   - Case studies
   - Specific illustrations

4. Key Findings & Implications:
   - Major conclusions
   - Practical applications
   - Significant insights

Rules:
- Use bullet points and sub-bullets for clear organization
- Keep points concise but informative
- Avoid paragraph-style writing
- Ensure logical grouping of related points
- Include specific details and examples where relevant

STRUCTURED SUMMARY:
"""
combine_prompt = PromptTemplate(template=combine_prompt_template, input_variables=["text"])

# QA Prompt can also be enhanced for more structured answers
qa_prompt_template = """
Use the following pieces of context to answer the question at the end. 
If you don't know the answer, just say that you don't know, don't try to make up an answer.

Context: {context}

Question: {question}

Please structure your answer in the following format when applicable:
1. Direct Answer:
   - Provide a clear, concise response

2. Supporting Details:
   - List relevant facts and information
   - Include specific examples from the context

3. Additional Context:
   - Any relevant background information
   - Related concepts or implications

4. Sources:
   - Reference specific parts of the context used

Present the information in a point-wise manner using appropriate bullet points and numbering.
Avoid long paragraphs unless absolutely necessary.
"""
qa_prompt = PromptTemplate(
    template=qa_prompt_template,
    input_variables=["context", "question"]
)

# Constants
CHUNK_SIZE = 500
CHUNK_OVERLAP = 50

# Utility Functions
def extract_text_from_document(file):
    """Extract text from uploaded document files."""
    if file.name.endswith(".pdf"):
        pdf_reader = PdfReader(file)
        return " ".join(page.extract_text() for page in pdf_reader.pages)
    elif file.name.endswith(".docx"):
        doc = docx.Document(file)
        return " ".join(paragraph.text for paragraph in doc.paragraphs)
    elif file.name.endswith(".txt"):
        return file.read().decode("utf-8")
    else:
        return ""

def process_content(content, source):
    """Process content into chunks and create Document objects."""
    words = content.split()
    chunks = []
    for i in range(0, len(words), CHUNK_SIZE):
        chunk = " ".join(words[i:i + CHUNK_SIZE + CHUNK_OVERLAP])
        chunks.append(Document(page_content=chunk, metadata={"source": source}))
    return chunks

def download_audio_ytdlp(video_url):
    """Download audio from YouTube using yt-dlp."""
    ydl_opts = {
        "format": "bestaudio/best",
        "outtmpl": tempfile.NamedTemporaryFile(delete=False, suffix=".mp3").name,
        "quiet": True,
    }
    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        info_dict = ydl.extract_info(video_url, download=True)
        return ydl.prepare_filename(info_dict)

def fetch_youtube_transcript_or_whisper(video_id):
    """Fetch transcript or generate using Whisper."""
    video_url = f"https://www.youtube.com/watch?v={video_id}"
    try:
        transcript = YouTubeTranscriptApi.get_transcript(video_id, languages=["en"])
        return " ".join([t["text"] for t in transcript])
    except (NoTranscriptFound, TranscriptsDisabled):
        st.warning("No transcript available. Attempting Whisper transcription...")
        try:
            audio_path = download_audio_ytdlp(video_url)
            model = whisper.load_model("base")
            transcription = model.transcribe(audio_path)
            os.remove(audio_path)
            return transcription["text"]
        except Exception as e:
            st.error(f"Failed to transcribe audio: {e}")
            return ""

# Initialize session states
if "vectorstore" not in st.session_state:
    st.session_state.vectorstore = None
if "current_summary" not in st.session_state:
    st.session_state.current_summary = ""
if "current_qa_pairs" not in st.session_state:
    st.session_state.current_qa_pairs = []
if "history" not in st.session_state:
    st.session_state.history = []

# Embedding initialization
embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

# Content Type Selection
st.subheader("Choose the type of content to summarize")
content_type = st.selectbox("Select content type", ["URL", "Audio File", "Document File"])
input_content = None

if content_type == "URL":
    input_content = st.text_input("Enter the URL to summarize")
elif content_type == "Audio File":
    input_content = st.file_uploader("Upload an audio file (mp3, wav, m4a)", type=["mp3", "wav", "m4a"])
elif content_type == "Document File":
    input_content = st.file_uploader("Upload a document file (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])

# Process Content Button
if st.button("Process Content"):
    if not groq_api_key.strip() or not input_content:
        st.error("Please provide valid input and Groq API Key.")
    else:
        try:
            with st.spinner("Processing content..."):
                docs = []
                source_name = ""

                # Get specialized model for content type
                model_type = "youtube" if content_type == "URL" and ("youtube.com" in input_content or "youtu.be" in input_content) else content_type.lower().split()[0]
                llm = ModelConfig().get_model(model_type, groq_api_key)

                # Process content based on type
                if content_type == "URL":
                    source_name = input_content
                    if "youtube.com" in input_content or "youtu.be" in input_content:
                        video_id = (
                            input_content.split("v=")[-1].split("&")[0]
                            if "youtube.com" in input_content
                            else input_content.split("/")[-1]
                        )
                        content = fetch_youtube_transcript_or_whisper(video_id)
                        if content:
                            docs.extend(process_content(content, input_content))
                        else:
                            st.error("Failed to process YouTube video. No content available.")
                    else:
                        try:
                            loader = UnstructuredURLLoader(urls=[input_content])
                            unstructured_docs = loader.load()
                            for doc in unstructured_docs:
                                docs.extend(process_content(doc.page_content, input_content))
                        except Exception as e:
                            st.error(f"Failed to load URL content: {e}")

                elif content_type == "Audio File":
                    source_name = input_content.name
                    with tempfile.NamedTemporaryFile(delete=True) as temp_file:
                        temp_file.write(input_content.read())
                        temp_file.flush()
                        model = whisper.load_model("base")
                        transcription = model.transcribe(temp_file.name)
                        audio_content = transcription["text"]
                        docs.extend(process_content(audio_content, "Uploaded Audio"))

                elif content_type == "Document File":
                    source_name = input_content.name
                    doc_content = extract_text_from_document(input_content)
                    docs.extend(process_content(doc_content, input_content.name))

                if not docs:
                    st.error("No content found. Please check the input and try again.")
                else:
                    st.session_state.vectorstore = FAISS.from_documents(docs, embeddings)

                    summarize_chain = load_summarize_chain(
                        llm,
                        chain_type="map_reduce",
                        map_prompt=map_prompt,
                        combine_prompt=combine_prompt,
                        verbose=True
                    )
                    st.session_state.current_summary = summarize_chain.run(docs)
                    st.session_state.current_qa_pairs = []

                    # Add to history
                    st.session_state.history.append({
                        "type": content_type,
                        "name": source_name,
                        "summary": st.session_state.current_summary,
                        "qa_pairs": []
                    })
                    
                    st.success("Content processed successfully!")

        except Exception as e:
            st.error(f"An error occurred: {e}")
            st.exception(e)

# Display Summary and Q&A
if st.session_state.current_summary:
    st.subheader("Summary")
    st.write(st.session_state.current_summary)

if st.session_state.vectorstore:
    st.subheader("Chat with the Data")
    user_question = st.text_input("Ask a question about the content:")
    
    if user_question:
        with st.spinner("Thinking..."):
            # Create Mixtral model instance directly for Q&A
            llm = ChatGroq(
                model="mixtral-8x7b-32768",
                groq_api_key=groq_api_key
            )
            
            qa_chain = RetrievalQA.from_chain_type(
                llm=llm,
                chain_type="stuff",
                retriever=st.session_state.vectorstore.as_retriever(
                    search_kwargs={"k": 3}
                ),
                chain_type_kwargs={"prompt": qa_prompt},
                return_source_documents=True,
            )
            
            response = qa_chain({"query": user_question})
            result = response["result"]
            source_documents = response["source_documents"]
            
            st.write("**Answer:**", result)
            
            with st.expander("Source Documents"):
                for doc in source_documents:
                    st.write(f"- {doc.page_content[:200]}...")
            
            qa_pair = {
                "question": user_question,
                "answer": result,
                "sources": [doc.metadata.get("source", "Unknown Source") for doc in source_documents]
            }
            st.session_state.current_qa_pairs.append(qa_pair)
            st.session_state.history[-1]["qa_pairs"].append(qa_pair)

# Update Sidebar History
with st.sidebar:
    st.subheader("History")
    if st.session_state.history:
        for idx, entry in enumerate(st.session_state.history):
            with st.expander(f"Session {idx + 1}"):
                st.write("**Summary:**", entry["summary"])
                
                if "qa_pairs" in entry and entry["qa_pairs"]:
                    st.write("**Q&A Pairs:**")
                    for qidx, qa in enumerate(entry["qa_pairs"]):
                        st.write(f"**Q{qidx + 1}:** {qa['question']}")
                        st.write(f"**A{qidx + 1}:** {qa['answer']}")
                        if "sources" in qa and qa["sources"]:
                            st.write("**Sources:**")
                            for source in qa["sources"]:
                                st.write(f"- {source}")
    else:
        st.write("No history available.")
