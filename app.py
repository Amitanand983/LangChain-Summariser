import validators
import streamlit as st
from langchain.prompts import PromptTemplate
from langchain_groq import ChatGroq
from langchain.chains import RetrievalQA
from langchain.schema import Document
from langchain.vectorstores import FAISS
from langchain.embeddings import HuggingFaceEmbeddings
from langchain_community.document_loaders import UnstructuredURLLoader
from youtube_transcript_api import YouTubeTranscriptApi, TranscriptsDisabled, NoTranscriptFound
from transformers import pipeline, AutoTokenizer
import whisper
from PyPDF2 import PdfReader
import docx
import tempfile
import os
import pickle

# Streamlit App Configuration
st.set_page_config(
    page_title="LangChain: Summarize and Chat",
    page_icon="🦜",
)
st.title("🦜 LangChain: Summarize and Chat")
st.subheader("Summarize URL, Audio, or Document and Chat with the Data")

# Sidebar for Groq API Key Input
with st.sidebar:
    groq_api_key = st.text_input("Groq API Key", value="", type="password")

# Inputs
generic_url = st.text_input("URL", label_visibility="collapsed")
audio_file = st.file_uploader("Upload an audio file (mp3, wav, m4a)", type=["mp3", "wav", "m4a"])
document_file = st.file_uploader("Upload a document file (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])

# LangChain LLM Setup (Gemma Model)
if groq_api_key.strip():
    llm = ChatGroq(model="Gemma-7b-It", groq_api_key=groq_api_key)
else:
    llm = None

# Prompt Template for Summarization
summarize_prompt_template = """
Summarize the following content concisely:
{text}
"""
summarize_prompt = PromptTemplate(template=summarize_prompt_template, input_variables=["text"])

# Token limits
MAX_CHUNK_TOKENS = 500  # Reduced to minimize latency

# Embedder for VectorStore
embeddings = HuggingFaceEmbeddings()

def split_into_chunks(text, max_tokens=MAX_CHUNK_TOKENS):
    """Split text into chunks of max_tokens length."""
    words = text.split()
    for i in range(0, len(words), max_tokens):
        yield " ".join(words[i:i + max_tokens])

def extract_text_from_document(file):
    """Extract text from uploaded document files."""
    if file.name.endswith(".pdf"):
        pdf_reader = PdfReader(file)
        return " ".join(page.extract_text() for page in pdf_reader.pages)
    elif file.name.endswith(".docx"):
        doc = docx.Document(file)
        return " ".join(paragraph.text for paragraph in doc.paragraphs)
    elif file.name.endswith(".txt"):
        return file.read().decode("utf-8")
    else:
        return ""

# Initialize summarizer model
summarizer = pipeline("summarization", model="t5-small", tokenizer="t5-small")  # Lightweight model

def summarize_chunks(chunks):
    """Summarize each chunk and combine results."""
    summaries = []
    for chunk in chunks:
        try:
            summary = summarizer(chunk, max_length=150, min_length=30, do_sample=False)
            summaries.append(summary[0]['summary_text'])
        except Exception as e:
            st.error(f"Failed to summarize chunk: {e}")
    return " ".join(summaries)

# Persistent storage for VectorStore
if "vectorstore" not in st.session_state:
    st.session_state.vectorstore = None
if "summarized_text" not in st.session_state:
    st.session_state.summarized_text = ""

# Summarization Button
if st.button("Summarize the Content"):
    if not llm or (not generic_url.strip() and not audio_file and not document_file):
        st.error("Please provide valid input and Groq API Key.")
    else:
        try:
            with st.spinner("Processing..."):
                docs = []

                # Process YouTube or Website URL
                if generic_url:
                    if "youtube.com" in generic_url or "youtu.be" in generic_url:
                        try:
                            video_id = (
                                generic_url.split("v=")[-1].split("&")[0]
                                if "youtube.com" in generic_url
                                else generic_url.split("/")[-1]
                            )
                            transcript = YouTubeTranscriptApi.get_transcript(video_id)
                            content = " ".join([t["text"] for t in transcript])
                            docs.append(Document(page_content=content, metadata={"source": generic_url}))
                        except TranscriptsDisabled:
                            st.error("Transcripts are disabled for this video. Please try another video.")
                        except Exception as e:
                            st.error(f"Failed to fetch YouTube content: {e}")
                    else:
                        loader = UnstructuredURLLoader(urls=[generic_url])
                        unstructured_docs = loader.load()
                        docs.extend(unstructured_docs)

                # Process Audio File
                if audio_file:
                    with tempfile.NamedTemporaryFile(delete=True) as temp_file:
                        temp_file.write(audio_file.read())
                        temp_file.flush()
                        model = whisper.load_model("tiny")
                        transcription = model.transcribe(temp_file.name)
                        audio_content = transcription["text"]
                        docs.append(Document(page_content=audio_content, metadata={"source": "Uploaded Audio"}))

                # Process Document File
                if document_file:
                    doc_content = extract_text_from_document(document_file)
                    docs.append(Document(page_content=doc_content, metadata={"source": document_file.name}))

                # Handle Empty or Invalid Data
                if not docs:
                    st.error("No content found. Please check the input and try again.")
                else:
                    all_summaries = []
                    for doc in docs:
                        chunks = list(split_into_chunks(doc.page_content))
                        summary = summarize_chunks(chunks)
                        all_summaries.append(summary)

                    # Store Summarized Content in VectorStore
                    st.session_state.summarized_text = "\n\n".join(all_summaries)
                    st.session_state.vectorstore = FAISS.from_texts([st.session_state.summarized_text], embeddings)
                    st.success(st.session_state.summarized_text)

        except Exception as e:
            st.error(f"An error occurred: {e}")

# Chatbot Interface
if st.session_state.vectorstore:
    st.subheader("Chat with the Data")
    user_question = st.text_input("Ask a question about the summarized data:")
    if user_question:
        with st.spinner("Thinking..."):
            qa_chain = RetrievalQA.from_chain_type(
                llm=llm,
                chain_type="stuff",
                retriever=st.session_state.vectorstore.as_retriever(k=5),  # Increase k for better answers
                return_source_documents=True
            )
            result = qa_chain({"query": user_question})
            answer = result["result"]
            sources = result.get("source_documents", [])

            # Display the answer
            st.success(answer)
            if sources:
                st.write("### Sources:")
                for source in sources:
                    st.write(f"- {source.metadata.get('source', 'Unknown Source')}")
                    
## Working version of chabot -1
# import validators
# import streamlit as st
# from langchain.prompts import PromptTemplate
# from langchain_groq import ChatGroq
# from langchain.chains.summarize import load_summarize_chain
# from langchain.chains import RetrievalQA
# from langchain.schema import Document
# from langchain.vectorstores import FAISS
# from langchain.embeddings import HuggingFaceEmbeddings
# from langchain_community.document_loaders import YoutubeLoader, UnstructuredURLLoader
# from youtube_transcript_api import YouTubeTranscriptApi
# from youtube_transcript_api._errors import TranscriptsDisabled, NoTranscriptFound
# import whisper
# from PyPDF2 import PdfReader
# import docx
# import tempfile
# import os
# import pickle

# # Streamlit App Configuration
# st.set_page_config(
#     page_title="LangChain: Summarize and Chat",
#     page_icon="🦜",
# )
# st.title("🦜 LangChain: Summarize and Chat")
# st.subheader("Summarize URL, Audio, or Document and Chat with the Data")

# # Sidebar for Groq API Key Input
# with st.sidebar:
#     groq_api_key = st.text_input("Groq API Key", value="", type="password")

# # Inputs
# generic_url = st.text_input("URL", label_visibility="collapsed")
# audio_file = st.file_uploader("Upload an audio file (mp3, wav, m4a)", type=["mp3", "wav", "m4a"])
# document_file = st.file_uploader("Upload a document file (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])

# # LangChain LLM Setup (Gemma Model)
# llm = ChatGroq(model="Gemma-7b-It", groq_api_key=groq_api_key)

# # Prompt Template for Summarization
# summarize_prompt_template = """
# Summarize the following content concisely:
# {text}
# """
# summarize_prompt = PromptTemplate(template=summarize_prompt_template, input_variables=["text"])

# # Token limits
# MAX_CHUNK_TOKENS = 1000  # Max tokens for each chunk, adjusted for safety
# TOKEN_RESERVE = 1000  # Reserve tokens for metadata and prompts

# # Embedder for VectorStore
# embeddings = HuggingFaceEmbeddings()

# def split_into_chunks(text, max_tokens):
#     """Split text into chunks that fit the token limit."""
#     words = text.split()
#     for i in range(0, len(words), max_tokens):
#         yield " ".join(words[i:i + max_tokens])

# def extract_text_from_document(file):
#     """Extract text from uploaded document files."""
#     if file.name.endswith(".pdf"):
#         pdf_reader = PdfReader(file)
#         return " ".join(page.extract_text() for page in pdf_reader.pages)
#     elif file.name.endswith(".docx"):
#         doc = docx.Document(file)
#         return " ".join(paragraph.text for paragraph in doc.paragraphs)
#     elif file.name.endswith(".txt"):
#         return file.read().decode("utf-8")
#     else:
#         return ""

# def summarize_chunks(llm, chunks):
#     """Summarize each chunk and combine results."""
#     summaries = []
#     summarize_chain = load_summarize_chain(llm, chain_type="stuff", prompt=summarize_prompt)
#     for chunk in chunks:
#         doc = Document(page_content=chunk, metadata={})
#         summary = summarize_chain.run([doc])
#         summaries.append(summary)
#     return " ".join(summaries)

# # Initialize persistent storage for VectorStore
# if "vectorstore" not in st.session_state:
#     st.session_state.vectorstore = None
# if "summarized_text" not in st.session_state:
#     st.session_state.summarized_text = ""

# # Button for Summarization
# if st.button("Summarize the Content"):
#     if not groq_api_key.strip() or (not generic_url.strip() and not audio_file and not document_file):
#         st.error("Please provide valid input and Groq API Key.")
#     else:
#         try:
#             with st.spinner("Processing..."):
#                 docs = []

#                 # Process YouTube or Website URL
#                 if generic_url:
#                     if "youtube.com" in generic_url or "youtu.be" in generic_url:
#                         try:
#                             video_id = (
#                                 generic_url.split("v=")[-1].split("&")[0]
#                                 if "youtube.com" in generic_url
#                                 else generic_url.split("/")[-1]
#                             )
#                             transcript = YouTubeTranscriptApi.get_transcript(video_id)
#                             content = " ".join([t["text"] for t in transcript])
#                             docs.append(Document(page_content=content, metadata={"source": generic_url}))
#                         except TranscriptsDisabled:
#                             st.error("Transcripts are disabled for this video. Please try another video.")
#                         except Exception as e:
#                             st.error(f"Failed to fetch YouTube content: {e}")
#                     else:
#                         loader = UnstructuredURLLoader(urls=[generic_url])
#                         unstructured_docs = loader.load()
#                         for doc in unstructured_docs:
#                             docs.append(doc)

#                 # Process Audio File
#                 if audio_file:
#                     with tempfile.NamedTemporaryFile(delete=True) as temp_file:
#                         temp_file.write(audio_file.read())
#                         temp_file.flush()
#                         model = whisper.load_model("base")
#                         transcription = model.transcribe(temp_file.name)
#                         audio_content = transcription["text"]
#                         docs.append(Document(page_content=audio_content, metadata={"source": "Uploaded Audio"}))

#                 # Process Document File
#                 if document_file:
#                     doc_content = extract_text_from_document(document_file)
#                     docs.append(Document(page_content=doc_content, metadata={"source": document_file.name}))

#                 # Handle Empty or Invalid Data
#                 if not docs:
#                     st.error("No content found. Please check the input and try again.")
#                 else:
#                     # Split and Summarize
#                     all_summaries = []
#                     for doc in docs:
#                         chunks = list(split_into_chunks(doc.page_content, MAX_CHUNK_TOKENS))
#                         summary = summarize_chunks(llm, chunks)
#                         all_summaries.append(summary)

#                     # Store Summarized Content in VectorStore
#                     st.session_state.summarized_text = "\n\n".join(all_summaries)
#                     st.session_state.vectorstore = FAISS.from_texts([st.session_state.summarized_text], embeddings)
#                     st.success(st.session_state.summarized_text)

#         except Exception as e:
#             st.error(f"An error occurred: {e}")
#             st.exception(e)

# # Chatbot Interface
# if st.session_state.vectorstore:
#     st.subheader("Chat with the Data")
#     user_question = st.text_input("Ask a question about the summarized data:")
#     if user_question:
#         with st.spinner("Thinking..."):
#             # Define RetrievalQA chain
#             qa_chain = RetrievalQA.from_chain_type(
#                 llm=llm,
#                 chain_type="stuff",
#                 retriever=st.session_state.vectorstore.as_retriever(),
#                 return_source_documents=True
#             )
#             # Get the answer
#             result = qa_chain({"query": user_question})
#             answer = result["result"]
#             sources = result.get("source_documents", [])

#             # Display the answer
#             st.success(answer)
#             if sources:
#                 st.write("### Sources:")
#                 for source in sources:
#                     st.write(f"- {source.metadata.get('source', 'Unknown Source')}")


### Working code till Audio, youtube with multiple language transcripts, and pdf, docx, txt

# import validators
# import tempfile
# import streamlit as st
# from langchain.prompts import PromptTemplate
# from langchain_groq import ChatGroq
# from langchain.chains.summarize import load_summarize_chain
# from langchain.schema import Document
# from langchain_community.document_loaders import UnstructuredURLLoader
# from youtube_transcript_api import YouTubeTranscriptApi
# from youtube_transcript_api._errors import TranscriptsDisabled, NoTranscriptFound
# import whisper
# from PyPDF2 import PdfReader
# import docx

# # Streamlit App Configuration
# st.set_page_config(
#     page_title="LangChain: Summarize Text From YT, Website, Audio, or Document",
#     page_icon="🦜",
# )
# st.title("🦜 LangChain: Summarize Text From YT, Website, Audio, or Document")
# st.subheader("Summarize URL, Audio, or Document")

# # Sidebar for Groq API Key Input
# with st.sidebar:
#     groq_api_key = st.text_input("Groq API Key", value="", type="password")

# # Inputs
# generic_url = st.text_input("URL", label_visibility="collapsed")
# audio_file = st.file_uploader("Upload an audio file (mp3, wav, m4a)", type=["mp3", "wav", "m4a"])
# document_file = st.file_uploader("Upload a document file (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])

# # LangChain LLM Setup (Gemma Model)
# llm = ChatGroq(model="Gemma-7b-It", groq_api_key=groq_api_key)

# # Prompt Template
# prompt_template = """
# Summarize the following content concisely:
# {text}
# """
# prompt = PromptTemplate(template=prompt_template, input_variables=["text"])

# # Token limits
# MAX_CHUNK_TOKENS = 1000  # Max tokens for each chunk, adjusted for safety
# TOKEN_RESERVE = 1000  # Reserve tokens for metadata and prompts

# def split_into_chunks(text, max_tokens):
#     """Split text into chunks that fit the token limit."""
#     words = text.split()
#     for i in range(0, len(words), max_tokens):
#         yield " ".join(words[i:i + max_tokens])

# def extract_text_from_document(file):
#     """Extract text from uploaded document files."""
#     if file.name.endswith(".pdf"):
#         pdf_reader = PdfReader(file)
#         return " ".join(page.extract_text() for page in pdf_reader.pages)
#     elif file.name.endswith(".docx"):
#         doc = docx.Document(file)
#         return " ".join(paragraph.text for paragraph in doc.paragraphs)
#     elif file.name.endswith(".txt"):
#         return file.read().decode("utf-8")
#     else:
#         return ""

# def summarize_chunks(llm, chunks):
#     """Summarize each chunk and combine results."""
#     summaries = []
#     summarize_chain = load_summarize_chain(llm, chain_type="stuff", prompt=prompt)
#     for chunk in chunks:
#         doc = Document(page_content=chunk, metadata={})
#         summary = summarize_chain.run([doc])
#         summaries.append(summary)
#     return " ".join(summaries)

# # Button for Summarization
# if st.button("Summarize the Content"):
#     if not groq_api_key.strip() or (not generic_url.strip() and not audio_file and not document_file):
#         st.error("Please provide valid input and Groq API Key.")
#     else:
#         try:
#             with st.spinner("Processing..."):
#                 docs = []

#                 # Process YouTube or Website URL
#                 if generic_url:
#                     if "youtube.com" in generic_url or "youtu.be" in generic_url:
#                         try:
#                             video_id = (
#                                 generic_url.split("v=")[-1].split("&")[0]
#                                 if "youtube.com" in generic_url
#                                 else generic_url.split("/")[-1]
#                             )
#                             # Fetch available transcripts
#                             available_transcripts = YouTubeTranscriptApi.list_transcripts(video_id)
#                             transcript = None
#                             for lang in available_transcripts:
#                                 try:
#                                     transcript = available_transcripts.find_transcript([lang.language_code])
#                                     break
#                                 except Exception:
#                                     continue
#                             if not transcript:
#                                 st.error("No usable transcript found for this video.")
#                                 st.stop()
#                             content = " ".join([t["text"] for t in transcript.fetch()])
#                             docs.append(Document(page_content=content, metadata={"source": generic_url}))
#                         except TranscriptsDisabled:
#                             st.error("Transcripts are disabled for this video. Please try another video.")
#                         except NoTranscriptFound:
#                             st.error("No transcripts are available for this video.")
#                         except Exception as e:
#                             st.error(f"Failed to fetch YouTube content: {e}")
#                     else:
#                         loader = UnstructuredURLLoader(urls=[generic_url])
#                         unstructured_docs = loader.load()
#                         for doc in unstructured_docs:
#                             docs.append(doc)

#                 # Process Audio File
#                 if audio_file:
#                     with tempfile.NamedTemporaryFile(delete=True) as temp_file:
#                         temp_file.write(audio_file.read())
#                         model = whisper.load_model("base")
#                         transcription = model.transcribe(temp_file.name)
#                         audio_content = transcription["text"]
#                         docs.append(Document(page_content=audio_content, metadata={"source": "Uploaded Audio"}))

#                 # Process Document File
#                 if document_file:
#                     doc_content = extract_text_from_document(document_file)
#                     docs.append(Document(page_content=doc_content, metadata={"source": document_file.name}))

#                 # Handle Empty or Invalid Data
#                 if not docs:
#                     st.error("No content found. Please check the input and try again.")
#                 else:
#                     # Split and Summarize
#                     all_summaries = []
#                     for doc in docs:
#                         chunks = list(split_into_chunks(doc.page_content, MAX_CHUNK_TOKENS))
#                         summary = summarize_chunks(llm, chunks)
#                         all_summaries.append(summary)

#                     final_summary = "\n\n".join(all_summaries)
#                     st.success(final_summary)
#         except Exception as e:
#             st.error(f"An error occurred: {e}")
#             st.exception(e)


### Working till audio, url, youtube, pdf, docx

# import validators
# import streamlit as st
# from langchain.prompts import PromptTemplate
# from langchain_groq import ChatGroq
# from langchain.chains.summarize import load_summarize_chain
# from langchain.schema import Document
# from langchain_community.document_loaders import YoutubeLoader, UnstructuredURLLoader
# from youtube_transcript_api import YouTubeTranscriptApi
# from youtube_transcript_api._errors import TranscriptsDisabled
# import whisper
# from PyPDF2 import PdfReader
# import docx
# import tempfile

# # Streamlit App Configuration
# st.set_page_config(
#     page_title="LangChain: Summarize Text From YT, Website, Audio, or Document",
#     page_icon="🦜",
# )
# st.title("🦜 LangChain: Summarize Text From YT, Website, Audio, or Document")
# st.subheader("Summarize URL, Audio, or Document")

# # Sidebar for Groq API Key Input
# with st.sidebar:
#     groq_api_key = st.text_input("Groq API Key", value="", type="password")

# # Inputs
# generic_url = st.text_input("URL", label_visibility="collapsed")
# audio_file = st.file_uploader("Upload an audio file (mp3, wav, m4a)", type=["mp3", "wav", "m4a"])
# document_file = st.file_uploader("Upload a document file (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])

# # LangChain LLM Setup (Gemma Model)
# llm = ChatGroq(model="Gemma-7b-It", groq_api_key=groq_api_key)

# # Prompt Template
# prompt_template = """
# Summarize the following content concisely:
# {text}
# """
# prompt = PromptTemplate(template=prompt_template, input_variables=["text"])

# # Token limits
# MAX_CHUNK_TOKENS = 1000  # Max tokens for each chunk, adjusted for safety
# TOKEN_RESERVE = 1000  # Reserve tokens for metadata and prompts

# def split_into_chunks(text, max_tokens):
#     """Split text into chunks that fit the token limit."""
#     words = text.split()
#     for i in range(0, len(words), max_tokens):
#         yield " ".join(words[i:i + max_tokens])

# def extract_text_from_document(file):
#     """Extract text from uploaded document files."""
#     if file.name.endswith(".pdf"):
#         pdf_reader = PdfReader(file)
#         return " ".join(page.extract_text() for page in pdf_reader.pages)
#     elif file.name.endswith(".docx"):
#         doc = docx.Document(file)
#         return " ".join(paragraph.text for paragraph in doc.paragraphs)
#     elif file.name.endswith(".txt"):
#         return file.read().decode("utf-8")
#     else:
#         return ""

# def summarize_chunks(llm, chunks):
#     """Summarize each chunk and combine results."""
#     summaries = []
#     summarize_chain = load_summarize_chain(llm, chain_type="stuff", prompt=prompt)
#     for chunk in chunks:
#         doc = Document(page_content=chunk, metadata={})
#         summary = summarize_chain.run([doc])
#         summaries.append(summary)
#     return " ".join(summaries)

# # Button for Summarization
# if st.button("Summarize the Content"):
#     if not groq_api_key.strip() or (not generic_url.strip() and not audio_file and not document_file):
#         st.error("Please provide valid input and Groq API Key.")
#     else:
#         try:
#             with st.spinner("Processing..."):
#                 docs = []

#                 # Process YouTube or Website URL
#                 if generic_url:
#                     if "youtube.com" in generic_url or "youtu.be" in generic_url:
#                         try:
#                             video_id = (
#                                 generic_url.split("v=")[-1].split("&")[0]
#                                 if "youtube.com" in generic_url
#                                 else generic_url.split("/")[-1]
#                             )
#                             transcript = YouTubeTranscriptApi.get_transcript(video_id)
#                             content = " ".join([t["text"] for t in transcript])
#                             docs.append(Document(page_content=content, metadata={"source": generic_url}))
#                         except TranscriptsDisabled:
#                             st.error("Transcripts are disabled for this video. Please try another video.")
#                         except Exception as e:
#                             st.error(f"Failed to fetch YouTube content: {e}")
#                     else:
#                         loader = UnstructuredURLLoader(urls=[generic_url])
#                         unstructured_docs = loader.load()
#                         for doc in unstructured_docs:
#                             docs.append(doc)

#                 # Process Audio File
#                 if audio_file:
#                     model = whisper.load_model("base")
#                     transcription = model.transcribe(audio_file.name)
#                     audio_content = transcription["text"]
#                     docs.append(Document(page_content=audio_content, metadata={"source": "Uploaded Audio"}))

#                 # Process Document File
#                 if document_file:
#                     doc_content = extract_text_from_document(document_file)
#                     docs.append(Document(page_content=doc_content, metadata={"source": document_file.name}))

#                 # Handle Empty or Invalid Data
#                 if not docs:
#                     st.error("No content found. Please check the input and try again.")
#                 else:
#                     # Split and Summarize
#                     all_summaries = []
#                     for doc in docs:
#                         chunks = list(split_into_chunks(doc.page_content, MAX_CHUNK_TOKENS))
#                         summary = summarize_chunks(llm, chunks)
#                         all_summaries.append(summary)

#                     final_summary = "\n\n".join(all_summaries)
#                     st.success(final_summary)
#         except Exception as e:
#             st.error(f"An error occurred: {e}")
#             st.exception(e)


